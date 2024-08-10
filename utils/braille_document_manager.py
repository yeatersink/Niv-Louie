from nicegui import app, events, ui
import io
import os
from pathlib import Path
import pandas as pd
import json
from docx import Document
import warnings
from utils.project import project


appdata_dir=os.getenv("LOCALAPPDATA")
niv_louie_app_data=os.path.join(appdata_dir,"Niv_Louie")
os.makedirs(niv_louie_app_data,exist_ok=True)


#The braille_numbers.json file is opened and read in to the braille_numbers_object variable
braille_numbers_file=open("utils/braille_to_numbers.json",encoding="utf8")
braille_numbers_object=json.load(braille_numbers_file)

#The braille_test_converter.json file is opened and read in to the braille_test_object variable
braille_test_file=open("utils/braille_test_converter.json",encoding="utf8")
braille_converter_object=json.load(braille_test_file)


class DocumentManager:
    def __init__(self):
        self.document_name=None
        self.document_contents=None
        self.document_projects_to_use=None
        self.document_list=None

    def update_document_name(self,e:events.ValueChangeEventArguments):
        self.document_name=e.value

    def handle_document_upload(self, e: events.UploadEventArguments):
        document_folder= os.path.join(niv_louie_app_data,"documents")
        if os.path.exists(document_folder) == False:
            os.makedirs(document_folder,exist_ok=True)
        if e.name.split(".")[-1]=="docx":
            document=Document(io.BytesIO(e.content.read()))
            document.save(os.path.join(document_folder,e.name))
            self.document_name=e.name
            self.document_contents=document
        elif e.name.split(".")[-1]=="txt":
            with open(os.path.join(document_folder,e.name),"w",encoding="utf-8") as file:
                file.write(        io.StringIO(e.content.read().decode("utf-8")))
            ui.notify("Text document to convert has been saved. ")


    def update_document_projects_to_use(self,e:events.ValueChangeEventArguments):
        self.document_projects_to_use=e.value


    def remove_document(self):
        document_path=os.path.join(niv_louie_app_data,"documents",self.document_name)
        if os.path.exists(document_path):
            os.remove(document_path)
            ui.notify("Document has been Removed!")


    def convert_document(self):
        document_folder= os.path.join(niv_louie_app_data,"documents")
        if self.document_name.split(".")[-1]=="docx":
            braille_document=Document()
            for paragraph in self.document_contents.paragraphs:
                new_paragraph=convert_text_to_braille(self.document_name,paragraph)
                braille_document.add_paragraph(new_paragraph)
            braille_document.save(os.path.join(document_folder,self.document_name.split(".")[0]+"-braille."+self.document_name.split(".")[-1]))
            ui.notify("Document converted.")
            ui.download(os.path.join(document_folder,self.document_name.split(".")[0]+"-braille."+self.document_name.split(".")[-1]))


def convert_text_to_braille(name,content):
    """
    This     function creates the braille document from the text document that the user has uploaded. It uses the braille_test_converter.json file to convert the text to braille. It also uses the braille_numbers.json file to convert numbers to braille. It also uses the language files that the user has chosen to convert the text to braille. It also checks if the text contains any characters that are not in the braille object and warns the user if it does.

    Parameters:
    (int): The index of the language that the user has chosen

    """

    #The language files are read in to pandas
    language_file_list=[]
    for selected_language in document.document_projects_to_use:
        language_file_list.append({"name":selected_language, "file":pd.read_csv(os.path.join(niv_louie_app_data,"languages","filtered_"+selected_language+".csv"),encoding="utf-8")})
    braille_content=""
    #This loop goes through each row in the content
    for row in content.text.split("\sn"):
        new_braille_content=row
        #This checks if the test contains any numbers
        if any(char.isdigit() for char in new_braille_content):
            new_text=""
            previous_was_number=False 
            #This loop goes through each character in the text
            for index,char in enumerate(new_braille_content):
                #This checks if the character is a number
                if char.isdigit() and previous_was_number==False:
                    #This adds a number sign to the character
                    new_text+="⠼"+char
                    previous_was_number=True
                else:
                    #This adds the character to the new text
                    new_text+=char
                    previous_was_number=False
            #This replaces the text with the new text
            new_braille_content=new_text
        for current_language in language_file_list:
            project.set_project_name(current_language["name"])
            project.set_all_fields()
            language_file=current_language["file"]
            #This loop goes through each character in the text
            for index,language_row in language_file.iterrows():
                if language_row[project.project_character_column] in new_braille_content and language_row[project.project_braille_column] != "nan":
                    new_braille_content=new_braille_content.replace(language_row[project.project_character_column],language_row[project. project_braille_column])
    #This loop goes through each character in the text and uses the braille test object to convert the text to braille
        for char in new_braille_content:
            if char in braille_converter_object:
                new_braille_content=new_braille_content.replace(char,braille_converter_object[char])
    #Checks if the test contains any non braille characters
        for char in new_braille_content:
            if char not in braille_numbers_object:
                warnings.warn("This document contains a character that is not in the braille object. This may be a mistake in your documen. Character: "+char)
        #This adds the new braille content to the braille content
        braille_content+=new_braille_content+"\n"
    return braille_content

document=DocumentManager()